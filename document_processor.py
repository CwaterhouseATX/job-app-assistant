"""Extract plain text from job posting URLs and local files (PDF, DOCX, TXT)."""

from __future__ import annotations

import re
from pathlib import Path
from typing import Union

import fitz  # PyMuPDF
import requests
from bs4 import BeautifulSoup
from docx import Document

PathLike = Union[str, Path]

_DEFAULT_TIMEOUT = 30
_USER_AGENT = (
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
    "(KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"
)


class DocumentExtractionError(Exception):
    """Raised when text cannot be extracted from a URL or file."""


class DocumentProcessor:
    """Fetch HTML job pages and parse PDF/DOCX/TXT into plain text."""

    def extract_from_url(self, url: str, *, timeout: float = _DEFAULT_TIMEOUT) -> str:
        """
        Download a URL and return visible text (scripts/styles removed).

        Raises DocumentExtractionError on network/HTTP failures or empty content.
        """
        url = (url or "").strip()
        if not url:
            raise DocumentExtractionError("URL is empty.")

        try:
            response = requests.get(
                url,
                timeout=timeout,
                headers={"User-Agent": _USER_AGENT},
                allow_redirects=True,
            )
            response.raise_for_status()
        except requests.RequestException as e:
            raise DocumentExtractionError(f"Failed to fetch URL: {e}") from e

        if not response.content:
            raise DocumentExtractionError("Empty response body.")

        # Prefer declared encoding; fall back to apparent_encoding for mislabeled pages.
        if response.encoding:
            try:
                html = response.text
            except UnicodeDecodeError:
                html = response.content.decode(
                    response.apparent_encoding or "utf-8", errors="replace"
                )
        else:
            html = response.content.decode(
                response.apparent_encoding or "utf-8", errors="replace"
            )

        text = self._html_to_text(html)
        text = self._normalize_whitespace(text)
        if not text:
            raise DocumentExtractionError("No extractable text from page HTML.")
        return text

    def extract_from_file(self, path: PathLike) -> str:
        """
        Extract text from a local file. Supported: .pdf, .docx, .txt

        Raises FileNotFoundError if the path does not exist.
        Raises DocumentExtractionError for unsupported types or empty extraction.
        """
        p = Path(path).expanduser().resolve()
        if not p.is_file():
            raise FileNotFoundError(f"Not a file: {p}")

        suffix = p.suffix.lower()
        if suffix == ".pdf":
            text = self._text_from_pdf(p)
        elif suffix == ".docx":
            text = self._text_from_docx(p)
        elif suffix == ".txt":
            text = self._text_from_txt(p)
        else:
            raise DocumentExtractionError(
                f"Unsupported file type {suffix!r}. Use .pdf, .docx, or .txt."
            )

        text = self._normalize_whitespace(text)
        if not text:
            raise DocumentExtractionError(f"No extractable text from file: {p}")
        return text

    @staticmethod
    def _html_to_text(html: str) -> str:
        soup = BeautifulSoup(html, "html.parser")
        for tag in soup(["script", "style", "noscript", "template"]):
            tag.decompose()
        raw = soup.get_text(separator="\n")
        return raw

    @staticmethod
    def _text_from_pdf(path: Path) -> str:
        try:
            doc = fitz.open(path)
        except Exception as e:
            raise DocumentExtractionError(f"Could not open PDF: {e}") from e
        try:
            parts: list[str] = []
            for page in doc:
                parts.append(page.get_text())
            return "\n".join(parts)
        finally:
            doc.close()

    @staticmethod
    def _text_from_docx(path: Path) -> str:
        try:
            document = Document(str(path))
        except Exception as e:
            raise DocumentExtractionError(f"Could not open DOCX: {e}") from e
        paragraphs = [para.text for para in document.paragraphs]
        # Tables are common in pasted job descriptions
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                paragraphs.append(" | ".join(cells))
        return "\n".join(paragraphs)

    @staticmethod
    def _text_from_txt(path: Path) -> str:
        data = path.read_bytes()
        for encoding in ("utf-8-sig", "utf-8", "cp1252", "latin-1"):
            try:
                return data.decode(encoding)
            except UnicodeDecodeError:
                continue
        return data.decode("utf-8", errors="replace")

    @staticmethod
    def _normalize_whitespace(text: str) -> str:
        text = text.replace("\r\n", "\n").replace("\r", "\n")
        text = re.sub(r"[ \t]+\n", "\n", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()
