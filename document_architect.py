"""Build ATS-friendly Word documents from AI-generated text (simple layout, standard fonts)."""

from __future__ import annotations

import re
from pathlib import Path
from typing import Iterable, Literal

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Pt

PathLike = str | Path

DEFAULT_FONT = "Calibri"
DEFAULT_BODY_PT = 11.0
DEFAULT_LINE_SPACING = 1.15


class DocumentArchitectError(Exception):
    """Raised when a document cannot be written."""


class DocumentArchitect:
    """
    Produces simple DOCX files suitable for ATS parsers: Calibri (or Arial), clear headings,
    bullets and paragraphs only — no images, text boxes, or complex tables.
    """

    def __init__(
        self,
        *,
        body_font: str = DEFAULT_FONT,
        body_size_pt: float = DEFAULT_BODY_PT,
        line_spacing: float = DEFAULT_LINE_SPACING,
    ) -> None:
        self.body_font = body_font
        self.body_size_pt = body_size_pt
        self.line_spacing = line_spacing

    def write_from_ai_text(
        self,
        text: str,
        output_path: PathLike,
        *,
        content_format: Literal["auto", "plain", "markdown"] = "auto",
    ) -> Path:
        """
        Turn model output into a .docx file.

        - ``plain``: split on blank lines into paragraphs.
        - ``markdown``: lightweight structure (#/##/### headings, ``-`` / ``*`` bullets,
          ``1.`` numbered lines). No tables or images.
        - ``auto``: use markdown mode if the text looks structured; otherwise plain.
        """
        raw = (text or "").strip()
        if not raw:
            raise DocumentArchitectError("Content is empty.")

        path = Path(output_path).expanduser().resolve()
        path.parent.mkdir(parents=True, exist_ok=True)

        mode = content_format
        if mode == "auto":
            mode = "markdown" if _looks_like_markdown_lite(raw) else "plain"

        doc = _new_document()
        self._apply_base_styles(doc)

        if mode == "plain":
            _emit_plain(doc, raw)
        else:
            _emit_markdown_lite(doc, raw)

        try:
            doc.save(str(path))
        except OSError as e:
            raise DocumentArchitectError(f"Could not save document: {e}") from e
        return path

    def write_sections(
        self,
        sections: Iterable[tuple[str, str]],
        output_path: PathLike,
        *,
        title: str | None = None,
    ) -> Path:
        """
        Write titled sections (e.g. Summary, Experience). Each section title uses Heading 2;
        body text uses plain paragraphs and optional inner markdown-lite lines.
        """
        path = Path(output_path).expanduser().resolve()
        path.parent.mkdir(parents=True, exist_ok=True)

        doc = _new_document()
        self._apply_base_styles(doc)

        if title:
            doc.add_heading(title.strip(), level=1)

        pairs = list(sections)
        if not pairs:
            raise DocumentArchitectError("No sections provided.")

        for heading, body in pairs:
            h = (heading or "").strip()
            b = (body or "").strip()
            if not h:
                continue
            doc.add_heading(h, level=2)
            if b:
                if _looks_like_markdown_lite(b):
                    _emit_markdown_lite(doc, b)
                else:
                    _emit_plain(doc, b)

        try:
            doc.save(str(path))
        except OSError as e:
            raise DocumentArchitectError(f"Could not save document: {e}") from e
        return path

    def _apply_base_styles(self, doc: Document) -> None:
        _style_normal_and_headings(
            doc,
            font_name=self.body_font,
            size_pt=self.body_size_pt,
            line_spacing=self.line_spacing,
        )


def _new_document() -> Document:
    return Document()


def _style_normal_and_headings(
    doc: Document,
    *,
    font_name: str,
    size_pt: float,
    line_spacing: float,
) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = font_name
    normal.font.size = Pt(size_pt)
    fmt = normal.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = line_spacing
    fmt.space_after = Pt(0)

    for name in ("Heading 1", "Heading 2", "Heading 3"):
        if name not in doc.styles:
            continue
        st = doc.styles[name]
        st.font.name = font_name
        if name == "Heading 1":
            st.font.size = Pt(max(size_pt + 6, 14))
        elif name == "Heading 2":
            st.font.size = Pt(max(size_pt + 3, 12))
        else:
            st.font.size = Pt(max(size_pt + 1, 11))
        st.font.bold = True

    for list_name in ("List Bullet", "List Number"):
        if list_name not in doc.styles:
            continue
        lp = doc.styles[list_name].paragraph_format
        lp.space_before = Pt(0)
        lp.space_after = Pt(0)
        lp.line_spacing_rule = WD_LINE_SPACING.SINGLE


def _apply_compact_list_paragraph_format(paragraph) -> None:
    pf = paragraph.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE


def _strip_markdown_artifacts(s: str) -> str:
    """Remove common inline Markdown markers without full parsing."""
    out = s
    while True:
        nxt = re.sub(r"\*\*([^*]+)\*\*", r"\1", out)
        if nxt == out:
            break
        out = nxt
    out = out.replace("**", "")
    out = re.sub(r"###\s*", "", out)
    out = re.sub(r"##\s*", "", out)
    return out


def _emit_plain(doc: Document, raw: str) -> None:
    blocks = [b.strip() for b in re.split(r"\n\s*\n+", raw) if b.strip()]
    if not blocks:
        inner = _strip_markdown_artifacts(raw.strip())
        if inner:
            doc.add_paragraph(inner)
        return
    for block in blocks:
        inner = re.sub(r"\s+", " ", block.replace("\n", " ")).strip()
        inner = _strip_markdown_artifacts(inner)
        if inner:
            doc.add_paragraph(inner)


def _emit_markdown_lite(doc: Document, raw: str) -> None:
    lines = raw.splitlines()
    buf: list[str] = []

    def flush_paragraph() -> None:
        if not buf:
            return
        text = re.sub(r"\s+", " ", " ".join(buf)).strip()
        buf.clear()
        text = _strip_markdown_artifacts(text)
        if text:
            doc.add_paragraph(text)

    for line in lines:
        stripped = line.strip()
        if not stripped:
            flush_paragraph()
            continue

        m = re.match(r"^(#{1,3})\s+(.*)$", stripped)
        if m:
            flush_paragraph()
            level = min(len(m.group(1)), 3)
            ht = _strip_markdown_artifacts(m.group(2).strip())
            if ht:
                doc.add_heading(ht, level=level)
            continue

        m = re.match(r"^[-*]\s+(.*)$", stripped)
        if m:
            flush_paragraph()
            bt = _strip_markdown_artifacts(m.group(1).strip())
            if bt:
                _apply_compact_list_paragraph_format(doc.add_paragraph(bt, style="List Bullet"))
            continue

        m = re.match(r"^\d+\.\s+(.*)$", stripped)
        if m:
            flush_paragraph()
            nt = _strip_markdown_artifacts(m.group(1).strip())
            if nt:
                _apply_compact_list_paragraph_format(doc.add_paragraph(nt, style="List Number"))
            continue

        buf.append(stripped)

    flush_paragraph()


def _looks_like_markdown_lite(text: str) -> bool:
    if re.search(r"(?m)^#{1,3}\s+\S", text):
        return True
    if re.search(r"(?m)^\s*[-*]\s+\S", text):
        return True
    if re.search(r"(?m)^\s*\d+\.\s+\S", text):
        return True
    return False
